from docx import Document as dc
import csv
import xlrd
import tkinter as tk

class editor(object):
    """read, edit, and write document"""

    def __init__(self):

        #set global variable
        pass

    def entryTemplateFieldContent(self, dir, csvData):

        #get this data from csv and loop
        readExcel = xlrd.open_workbook(csvData)
        sheet = readExcel.sheet_by_index(0)

        #iterate through data entry excel rows
        for x in range(sheet.nrows - 1):

            #read document
            doc = dc(dir)

            fileName = sheet.cell_value(x + 1,0)
            print(fileName)

            #iterate through data entry excel columns
            for cols in range(sheet.ncols):

                #read specific data from csv
                replaceFieldFound = sheet.cell_value(0, cols) #this value holds
                textReplacement = sheet.cell_value(x + 1, cols) #this value varies

                #for debug only
                #print(replaceFieldFound)
                #print(textReplacement)

                #read all paragraph (mailmerge mode)
                for paragraph in doc.paragraphs:

                    #read all runs
                    for run in paragraph.runs:

                        #replace field with desired text
                        if replaceFieldFound in run.text:

                            #notification
                            #"paragraph",paragraph,"run",run,
                            print("[",replaceFieldFound,"terganti dengan", textReplacement,"]")
                            run.text = run.text.replace(replaceFieldFound, str(textReplacement))

                            pass

                        pass

                    pass

                pass

            #create new file
            doc.save(fileName + ".docx")
            print("done! created file as",fileName + ".docx")
            #some bug casues print fist row only ! (Solved!)
            pass

        pass

    def entryTemplateFieldContentTable(self, dir, csvData, isAll, start, finish):

        #get this data from csv and loop
        readExcel = xlrd.open_workbook(csvData)
        sheet = readExcel.sheet_by_index(0)

        if isAll == True:

            rangeBegin = 0
            rangeEnd = sheet.nrows - 1

            pass

        else:

            rangeBegin = start - 1
            rangeEnd = finish

            pass

        #iterate through data entry excel rows
        for x in range(rangeBegin, rangeEnd):

            #read document
            doc = dc(dir)

            fileName = sheet.cell_value(x + 1,0)
            print(fileName)

            #iterate through data entry excel columns
            for cols in range(sheet.ncols):

                #read specific data from csv
                replaceFieldFound = sheet.cell_value(0, cols) #this value holds
                textReplacement = sheet.cell_value(x + 1, cols) #this value varies

                #for debug only
                #print(replaceFieldFound)
                #print(textReplacement)

                #replace table (mailmerge table data)
                for table in doc.tables:

                    for row in table.rows:

                        for cell in row.cells:

                            #read all paragraph (mailmerge mode)
                            for paragraph in cell.paragraphs:

                                #read all runs
                                for run in paragraph.runs:

                                    #print(run.text)

                                    #replace field with desired text
                                    if replaceFieldFound in run.text:

                                        #notification
                                        #"paragraph",paragraph,"run",run,
                                        print("[",replaceFieldFound,"terganti dengan", textReplacement,"]")
                                        run.text = run.text.replace(replaceFieldFound, str(textReplacement))

                                        pass

                                    pass

                                pass

                            pass

                        pass

                    pass

                #read all paragraph (mailmerge mode)
                for paragraph in doc.paragraphs:

                    #read all runs
                    for run in paragraph.runs:

                        #replace field with desired text
                        if replaceFieldFound in run.text:

                            #notification
                            #"paragraph",paragraph,"run",run,
                            print("[",replaceFieldFound,"terganti dengan", textReplacement,"]")
                            run.text = run.text.replace(replaceFieldFound, str(textReplacement))

                            pass

                        pass

                    pass

                pass

            #create new file
            doc.save(fileName + ".docx")
            print("done! created file as",fileName + ".docx")
            #some bug casues print fist row only ! (Solved!)
            pass

        pass


#create gui

windowEntry = tk.Tk()
windowTitle = windowEntry.title("Batch format generator! (~^3^)~")

windowWidth = 400
windowHeight = 300

#create initial canvas
canvas = tk.Canvas(windowEntry, width=windowWidth, height=windowHeight)
canvas.pack()

#label entry for word dir
wordText = tk.Label(windowEntry, text = "masukkan directory word: ")
canvas.create_window(windowWidth/2, 40, window=wordText)

#set word entry object
entryWord = tk.Entry(windowEntry)
canvas.create_window(windowWidth/2, 60, window=entryWord)

#label entry for excel dir
excelText = tk.Label(windowEntry, text = "masukkan directory excel: ")
canvas.create_window(windowWidth/2, 80, window=excelText)

#set excel entry object
entryExcel = tk.Entry(windowEntry)
canvas.create_window(windowWidth/2, 100, window=entryExcel)

#label entry for start row
startText = tk.Label(windowEntry, text = "masukan urutan row awal yang mau dibuat: ")
canvas.create_window(windowWidth/2, 120, window=startText)

#set start row
entryStartText = tk.Entry(windowEntry)
canvas.create_window(windowWidth/2, 140, window=entryStartText)

#label entry for stop row
stopText = tk.Label(windowEntry, text = "masukan urutan row akhir yang mau dibuat: ")
canvas.create_window(windowWidth/2, 160, window=stopText)

#set stop row
entryStopText = tk.Entry(windowEntry)
canvas.create_window(windowWidth/2, 180, window=entryStopText)

#excecute button



def executeClass():

    word = entryWord.get()
    excel = entryExcel.get()
    start = float(entryStartText.get())
    stop = float(entryStopText.get())
    file = editor().entryTemplateFieldContentTable(word, excel, False, int(start), int(stop))

    print(word)
    pass

button = tk.Button(text="GO!", command=executeClass)
canvas.create_window(windowWidth/2, 220, window=button)

windowEntry.mainloop()
