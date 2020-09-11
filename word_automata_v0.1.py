from docx import Document as dc
import csv
import xlrd

class editor(object):
    """read, edit, and write document"""

    def __init__(self):

        #set global variable
        pass

    def entryTemplateFieldContent(self, dir, csvData):

        #get this data from csv and loop
        readExcel = xlrd.open_workbook(csvData)
        sheet = readExcel.sheet_by_index(0)

        #iterate through data entry excel rows
        for x in range(sheet.nrows - 1):

            #read document
            doc = dc(dir)

            fileName = sheet.cell_value(x + 1,0)
            print(fileName)

            #iterate through data entry excel columns
            for cols in range(sheet.ncols):

                #read specific data from csv
                replaceFieldFound = sheet.cell_value(0, cols) #this value holds
                textReplacement = sheet.cell_value(x + 1, cols) #this value varies

                #for debug only
                #print(replaceFieldFound)
                #print(textReplacement)

                #read all paragraph (mailmerge mode)
                for paragraph in doc.paragraphs:

                    #read all runs
                    for run in paragraph.runs:

                        #replace field with desired text
                        if replaceFieldFound in run.text:

                            #notification
                            #"paragraph",paragraph,"run",run,
                            print("[",replaceFieldFound,"terganti dengan", textReplacement,"]")
                            run.text = run.text.replace(replaceFieldFound, str(textReplacement))

                            pass

                        pass

                    pass

                pass

            #create new file
            doc.save(fileName + ".docx")
            print("done! created file as",fileName + ".docx")
            #some bug casues print fist row only ! (Solved!)
            pass

        pass

    def entryTemplateFieldContentTable(self, dir, csvData):

        #get this data from csv and loop
        readExcel = xlrd.open_workbook(csvData)
        sheet = readExcel.sheet_by_index(0)

        #iterate through data entry excel rows
        for x in range(sheet.nrows - 1):

            #read document
            doc = dc(dir)

            fileName = sheet.cell_value(x + 1,0)
            print(fileName)

            #iterate through data entry excel columns
            for cols in range(sheet.ncols):

                #read specific data from csv
                replaceFieldFound = sheet.cell_value(0, cols) #this value holds
                textReplacement = sheet.cell_value(x + 1, cols) #this value varies

                #for debug only
                #print(replaceFieldFound)
                #print(textReplacement)

                #replace table (mailmerge table data)
                for table in doc.tables:

                    for row in table.rows:

                        for cell in row.cells:

                            #read all paragraph (mailmerge mode)
                            for paragraph in cell.paragraphs:

                                #read all runs
                                for run in paragraph.runs:

                                    #print(run.text)

                                    #replace field with desired text
                                    if replaceFieldFound in run.text:

                                        #notification
                                        #"paragraph",paragraph,"run",run,
                                        print("[",replaceFieldFound,"terganti dengan", textReplacement,"]")
                                        run.text = run.text.replace(replaceFieldFound, str(textReplacement))

                                        pass

                                    pass

                                pass

                            pass

                        pass

                    pass

                #read all paragraph (mailmerge mode)
                for paragraph in doc.paragraphs:

                    #read all runs
                    for run in paragraph.runs:

                        #replace field with desired text
                        if replaceFieldFound in run.text:

                            #notification
                            #"paragraph",paragraph,"run",run,
                            print("[",replaceFieldFound,"terganti dengan", textReplacement,"]")
                            run.text = run.text.replace(replaceFieldFound, str(textReplacement))

                            pass

                        pass

                    pass

                pass

            #create new file
            doc.save(fileName + ".docx")
            print("done! created file as",fileName + ".docx")
            #some bug casues print fist row only ! (Solved!)
            pass

        pass

while True:

    try:

        #document directory
        data = input("masukkan directory word: ")
        replace = input("masukan directory excel pengganti: ")

        #break ethernal loop
        if data == "quit":

            break

        #output = input("nama file output: ")

        file = editor().entryTemplateFieldContentTable(data, replace)

        pass

    except:

        #spit error
        print("error salah entry")

        pass
